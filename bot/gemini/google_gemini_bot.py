"""
Google gemini bot

@author zhayujie
@Date 2023/12/15
"""
# encoding:utf-8

import base64
import json
import os
import time
import re
import docx
from io import BytesIO
from pypdf import PdfReader
from bot.bot import Bot
import pandas as pd
import matplotlib.pyplot as plt
import google.generativeai as generativeai
from google.ai.generativelanguage_v1beta.types import content
from google import genai
from google.genai import types
from google.genai.types import Tool,GenerateContentConfig,GoogleSearch,Part,FunctionDeclaration,Type,FileData
from bot.session_manager import SessionManager
from bridge.context import ContextType, Context
from bridge.reply import Reply, ReplyType
from common.log import logger
from common.tmp_dir import TmpDir
from common import const, memory
from config import conf
from bot.baidu.baidu_wenxin_session import BaiduWenxinSession
from bot.gemini.google_genimi_vision import GeminiVision
from PIL import Image
from channel.telegram.telegram_text_util import escape
from channel.telegram.telegram_channel import TelegramChannel


# OpenAI对话模型API (可用)
class GoogleGeminiBot(Bot,GeminiVision):

    def __init__(self):
        super().__init__()
        # Add these imports at the top
        import asyncio
        import nest_asyncio
        
        # Initialize event loop
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        
        # Allow nested event loops
        nest_asyncio.apply()
        
        self.api_key = conf().get("gemini_api_key")
        self.model = conf().get('model')
        self.Model_ID = self.model.upper()
        self.system_prompt = conf().get("character_desc")
        self.function_call_dicts = {
            "screenplay_scenes_breakdown": self.screenplay_scenes_breakdown,
            "screenplay_assets_breakdown": self.screenplay_assets_breakdown
        }
        # 复用文心的token计算方式
        self.sessions = SessionManager(BaiduWenxinSession, model=self.model or "gpt-3.5-turbo")
        
        # Initialize a client according to old generativeai SDK
        generativeai.configure(api_key=self.api_key,transport='rest')
        self.generation_config = {
            "temperature": 0.4,
            "top_p": 0.95,
            "top_k": 64,
            "max_output_tokens": 65536,
        }
        self.safety_settings = [
            {
                "category": "HARM_CATEGORY_HARASSMENT",
                "threshold": "BLOCK_NONE"
            },
            {
                "category": "HARM_CATEGORY_HATE_SPEECH",
                "threshold": "BLOCK_NONE"
            },
            {
                "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                "threshold": "BLOCK_NONE"
            },
            {
                "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                "threshold": "BLOCK_NONE"
            },
        ]
        self.tool_config={'function_calling_config': 'AUTO'}

        self.generative_model = generativeai.GenerativeModel(
                        model_name=self.model,
                        generation_config=self.generation_config,
                        safety_settings=self.safety_settings,
                        system_instruction=self.system_prompt,
                        tools = [
                            generativeai.protos.Tool(
                                function_declarations = [
                                    generativeai.protos.FunctionDeclaration(
                                        name = "screenplay_scenes_breakdown",
                                        description = "第一步先拆解剧本场景，提取场号、场景、内外、日夜等基础信息，形成场景列表；第二步根据输入的剧本名称找到剧本文件，准确统计剧本总页数，总字数，每场戏的字数，补充到场景列表中。",
                                        parameters = content.Schema(
                                            type = content.Type.OBJECT,
                                            enum = [],
                                            required = ["screenplay_title","scenes_list"],
                                            properties = {
                                                "screenplay_title": content.Schema(
                                                    type = content.Type.STRING,
                                                    description = "剧本名称",
                                                ),
                                                "scenes_list": content.Schema(
                                                    type = content.Type.ARRAY,
                                                    items = content.Schema(
                                                        type = content.Type.OBJECT,
                                                        properties = {
                                                            "id": content.Schema(
                                                                type = content.Type.INTEGER,
                                                                description = "场号",
                                                            ),
                                                            "location": content.Schema(
                                                                type = content.Type.STRING,
                                                                description = "场景名称",
                                                            ),
                                                            "daynight": content.Schema(
                                                                type = content.Type.STRING,
                                                                description = "日景还是夜景",
                                                                enum = ["日","夜"]
                                                            ),
                                                            "envirement": content.Schema(
                                                                type = content.Type.STRING,
                                                                description = "室内环境还是室外环境",
                                                                enum = ["内","外"]
                                                            ),
                                                        },
                                                    ),
                                                ),
                                            },
                                        ),
                                    ),
                                ],
                            ),
                        ],
                        tool_config=self.tool_config
                    )

        # Initialize a client according to new genai SDK
        self.client = genai.Client(api_key=self.api_key)

         # schema for screenplay_scenes_breakdown need to be updated
        self.screenplay_scenes_breakdown_schema = FunctionDeclaration(
            name="screenplay_scenes_breakdown",
            description="拆解剧本一共两个步骤，这是第一步，拆解场景。阅读剧本，结合内容，先提取剧本名称和剧本总页数，再逐个拆解场景，提取场号、场景名称、内外、日夜等基础信息，做成场景列表",
            parameters=types.Schema(
                type=Type.OBJECT,
                required=["screenplay_title","scenes_list"],
                properties={
                    "screenplay_title": types.Schema(
                        type=Type.STRING,
                        description="从文档中提取出的剧本名称",
                    ),
                    "total_pages": types.Schema(
                        type=Type.INTEGER,
                        description="剧本总页数",
                    ),
                    "total_words": types.Schema(
                        type=Type.INTEGER,
                        description="剧本总字数",
                    ),
                    "scenes_list": types.Schema(
                        type=Type.ARRAY,
                        items=types.Schema(
                            type=Type.OBJECT,
                            properties={
                                "scene_id": types.Schema(
                                    type=Type.INTEGER,
                                    description="场号，包含正戏与彩蛋，彩蛋场号顺接正戏，比如正戏最后一场是95，彩蛋第一场就是96",
                                ),
                                "location": types.Schema(
                                    type=Type.STRING,
                                    description="场景名称，包含正戏与彩蛋",
                                ),
                                "daynight": types.Schema(
                                    type=Type.STRING,
                                    description="日景还是夜景，如果剧本中某一场次开头没有明确标注，根据该场次内容自行推断，并从枚举列表中提取相应推断结果",
                                    enum = ["日", "夜", "极昼", "极夜", "晨", "昏", "晨(极昼)", "夜(极昼)", "晨(极夜)", "夜(极夜)"],
                                ),
                                "envirement": types.Schema(
                                    type=Type.STRING,
                                    description="室内环境还是室外环境，如果剧本中某一场次开头没有明确标注，根据该场次内容自行推断，并从枚举列表中提取相应推断结果",
                                    enum = ["内", "外", "外/内", "内/外"],
                                ),
                                "words": types.Schema(
                                    type=Type.INTEGER,
                                    description="场次字数",
                                ),
                                "pages": types.Schema(
                                    type=Type.NUMBER,
                                    description="场次页数，计算方式为场次字数除以页均字数即(words / (total_words / total_pages))",
                                ),
                                "estimated_duration": types.Schema(
                                    type=Type.NUMBER,
                                    description="场次预计时长，单位是分钟，计算方式为场次页数的1.2倍即(pages * 1.2)",
                                ),
                                "assets_id": types.Schema(
                                    type=Type.ARRAY,
                                    description="场次引用资产的id，从screenplay_assets_breakdown的返回值中取",
                                    items=types.Schema(
                                        type = Type.STRING
                                    )
                                ),
                            },
                        ),
                    ),
                },
            ),
        )
        self.screenplay_assets_breakdown_schema = FunctionDeclaration(
            name="screenplay_assets_breakdown",
            description="拆解剧本一共分两个步骤，这是第二步，拆解资产。阅读剧本，熟悉内容，逐场提取场景、人物和道具名称，制作成资产列表。该工具需要跟sreenplay_scenes_breakdown同时平行调用，才能完成剧本拆解。",
            parameters=types.Schema(
                type=Type.OBJECT,
                required=["screenplay_title","assets_list"],
                properties={
                    "screenplay_title": types.Schema(
                        type=Type.STRING,
                        description="剧本名称，调用函数必需输入的参数，仔细阅读文档，提取出剧本名称",
                    ),
                    "assets_list": types.Schema(
                        type=Type.ARRAY,
                        items=types.Schema(
                            type=Type.OBJECT,
                            properties={
                                "asset_type": types.Schema(
                                    type = types.Type.STRING,
                                    description = "资产类型，cast是有名字、有台词、对剧情发展有重要影响的角色；stunt是特技或动作演员，extra是没名字、没台词的群众演员；costume是服装，location是不需要陈设的外景，set是需要陈设的内景，vehicle是交通工具，animal是需要驯养员看守的动物，prop是道具",
                                    enum = [
                                    "cast",
                                    "stunt",
                                    "extra",
                                    "costume",
                                    "location",
                                    "set",
                                    "vehicle",
                                    "animal",
                                    "prop"
                                    ],
                                ),
                                "asset_id": types.Schema(
                                    type=Type.STRING,
                                    description="资产ID，格式为资产类型+两位阿拉伯数字即asset_type##，如cast01，location02，prop03, vehicle04等",
                                ),
                                "name": types.Schema(
                                    type=Type.STRING,
                                    description="资产名称,完整保留每一场中出现的资产，包括无名无姓的角色、一闪而过的场景、无足轻重的道具，不要简化或过滤。",
                                ),
                                "scene_ids": types.Schema(
                                    type=Type.ARRAY,
                                    description="资产所出现的场次",
                                    items=types.Schema(
                                        type = Type.INTEGER
                                    )
                                ),
                                "asset_pages": types.Schema(
                                    type=Type.NUMBER,
                                    description="资产所出现场次总页数",
                                ),
                                "estimated_asset_duration": types.Schema(
                                    type=Type.NUMBER,
                                    description="预计资产所出现场次总时长，单位是分钟，计算方式为资产页数的1.2倍即(asset_pages * 1.2)",
                                )
                            },
                        ),
                    ),
                },
            ),
        )
        self.function_declarations = Tool(function_declarations=[self.screenplay_scenes_breakdown_schema, self.screenplay_assets_breakdown_schema])
        self.google_search_tool = Tool(google_search=GoogleSearch())
        self.chat = self.client.chats.create(
            model=self.model,
            config=GenerateContentConfig(
                system_instruction=self.system_prompt,
                safety_settings=self.safety_settings,
                tools=[self.function_declarations],
                tool_config={
                    'function_calling_config': {
                        'mode': 'AUTO' 
                    }
                },
                response_modalities=['TEXT'],
                **self.generation_config
            )
        )
        self.image_chat = self.client.chats.create(
            model='gemini-2.0-flash-exp',
            config=GenerateContentConfig(
                safety_settings=self.safety_settings,
                response_modalities=['TEXT','Image'],
                **self.generation_config
            )
        )
        self.search_config = GenerateContentConfig(
            system_instruction=self.system_prompt,
            safety_settings=self.safety_settings,
            tools=[self.google_search_tool],
            response_modalities=['TEXT'],
            **self.generation_config
        )

    def reply(self, query, context: Context = None) -> Reply:
        try:
            if context.type == ContextType.FILE:
                mime_type = context.content[(context.content.rindex('.') + 1):]
                if mime_type in const.AUDIO or mime_type in const.VIDEO or mime_type in const.DOCUMENT or mime_type in const.TXT:
                    session_id = context["session_id"]
                    session = self.sessions.session_query(query, session_id)
                    return self.gemini_15_media(query, context, session)
                elif mime_type in const.DOCUMENT:
                    self._file_cache(context)
                    doc_cache = memory.USER_FILE_CACHE.get(context['session_id'])
                    return self._file_download(doc_cache)
            elif context.type == ContextType.IMAGE:
                if (self.model in const.GEMINI_15_FLASH_LIST or 
                    self.model in const.GEMINI_15_PRO_LIST or 
                    self.model in const.GEMINI_2_FLASH_LIST or 
                    self.model in const.GEMINI_25_PRO_LIST):
                    session_id = context["session_id"]
                    session = self.sessions.session_query(query, session_id)
                    return self.gemini_15_media(query, context, session)
                elif self.model in const.GEMINI_1_PRO_LIST:
                    memory.USER_IMAGE_CACHE[context["session_id"]] = {
                    "path": context.content,
                    "msg": context.get("msg")
                    }
                    logger.info(f"{context.content} cached to memory")
                    return None
            elif context.type == ContextType.VIDEO:
                session_id = context["session_id"]
                session = self.sessions.session_query(query, session_id)
                return self.gemini_15_media(query, context, session)
            elif context.type == ContextType.SHARING:
                session_id = context["session_id"]
                session = self.sessions.session_query(query, session_id)
                return self.gemini_15_media(query, context, session)
            elif context.type == ContextType.TEXT:
                logger.info(f"[{self.Model_ID}] query={query}")
                session_id = context["session_id"]
                session = self.sessions.session_query(query, session_id)

                # Set up the model
                if self.model in const.GEMINI_1_PRO_LIST:
                    gemini_messages = self._convert_to_gemini_1_messages(self._filter_gemini_1_messages(session.messages))
                    system_prompt = None
                    vision_res = self.do_vision_completion_if_need(session_id,query) # Image recongnition and vision completion
                    if vision_res:
                        return vision_res
                elif (self.model in const.GEMINI_15_PRO_LIST or 
                      self.model in const.GEMINI_15_FLASH_LIST or 
                      self.model in const.GEMINI_2_FLASH_LIST or 
                      self.model in const.GEMINI_25_PRO_LIST):
                    gemini_messages = self._convert_to_gemini_15_messages(session.messages)

                if self.model in const.GEMINI_GENAI_SDK:
                    # 检查缓存中是否媒体文件
                    file_cache = memory.USER_IMAGE_CACHE.get(session_id)
                    if file_cache:
                        first_data = file_cache['files'][0]
                        data_type = type(first_data).__name__
                        if data_type == 'dict':
                            mime_type = first_data.get('mime_type')
                            if mime_type in ['application/docx', 'application/doc']:
                                query = query + '\n' + first_data.get('data')
                            elif mime_type == 'application/pdf':
                                file_content = [Part.from_bytes(**first_data)]
                                file_content.append(query)
                                query = file_content
                        elif data_type in ['JpegImageFile','File']:
                            file_cache['files'].append(query)
                            query = file_cache['files']
                        elif data_type in ['FileData']:
                            filedata = [
                                {
                                'fileData': first_data
                            }
                            ]
                            filedata.append(query)
                            query = filedata
                        memory.USER_IMAGE_CACHE.pop(session_id)
                    if TelegramChannel().searching is True:
                        response = self.chat.send_message(query,config=self.search_config)
                    elif TelegramChannel().searching is False:
                        if TelegramChannel().imaging is True:
                            response = self.image_chat.send_message(query)
                        elif TelegramChannel().imaging is False:
                            response = self.chat.send_message(query)

                elif self.model not in const.GEMINI_GENAI_SDK:
                    chat_session = self.generative_model.start_chat(
                        # 消息堆栈中的最新数据抛出去，留给send_message方法，从query中拿
                        history=gemini_messages[:-1],
                        enable_automatic_function_calling=True
                    )
                    response = chat_session.send_message(query)
                
                # 轮询模型响应结果中是否有函数调用
                function_calling = True
                while function_calling:
                    if hasattr(response, 'function_calls'):
                        function_calls = response.function_calls if response.function_calls is not None else []
                    else:
                        function_calls = response.parts if response.parts[0].function_call.args is not None else []
                    function_response_parts = []
                    for part in function_calls:
                        function_call_reply = self.function_call_reply(part)
                        fn_name = function_call_reply.get('functionCall').get('name')
                        fn_args = function_call_reply.get('functionCall').get('args')
                        # 将reply_text转换为字符串
                        function_call_str = json.dumps(function_call_reply)
                        # add function call to session as model/assistant message
                        self.sessions.session_reply(function_call_str, session_id)
                        # call function
                        function_call = self.function_call_dicts.get(fn_name)
                        # 从fn_args中获取function_call的参数
                        function_response_part, function_response_str = function_call(context, fn_name, **fn_args)
                        function_response_parts.extend(function_response_part)
                        # add function response to session as user message
                        self.sessions.session_query(function_response_str, session_id)

                    if function_calls:
                        # new turn of model request with function response
                        if self.model in const.GEMINI_GENAI_SDK:
                            response = self.chat.send_message(function_response_parts)
                        else:
                            gemini_messages = self._convert_to_gemini_15_messages(session.messages)
                            chat_session = self.generative_model.start_chat(
                            # 消息堆栈中的最新数据抛出去，留给send_message方法，从function_response_str中拿
                                history=gemini_messages[:-1],
                                enable_automatic_function_calling=True
                            )
                            response = chat_session.send_message(function_response_part)
                        continue
                    function_calling = False

                if TelegramChannel().imaging is True:
                    self.get_reply_images(context, response)
                    return None
                elif TelegramChannel().imaging is False:
                    raw_reply_contents = self.get_reply_text(response)
                    reply_text = escape(raw_reply_contents)

                    # attach search sources from the Grounded response
                    if response.candidates[0].grounding_metadata:
                        grounding_metadata = response.candidates[0].grounding_metadata.grounding_chunks
                        if grounding_metadata:
                            inline_url = self.get_search_sources(response)
                            reply_text = f'{reply_text}\n\n{inline_url}'

                    self.sessions.session_reply(reply_text, session_id)
                    logger.info(f"[{self.Model_ID}] reply={reply_text}")
                    return Reply(ReplyType.TEXT, reply_text)
            else:
                logger.warning(f"[{self.Model_ID}] Unsupported message type, type={context.type}")
                return Reply(ReplyType.ERROR, f"[{self.Model_ID}] Unsupported message type, type={context.type}")
        except Exception as e:
            logger.error("[{}] fetch reply error, {}".format(self.Model_ID, e))
            return Reply(ReplyType.ERROR, f"[{self.Model_ID}] {e}")

    def function_call_reply(self, part):
        if hasattr(part, 'function_call'):
            fn = part.function_call
            fn_dict = type(fn).to_dict(fn)
            fn_name = fn_dict.get('name')
            fn_args = fn_dict.get('args')
        else:
            fn_name = part.name
            if "assets" in fn_name.lower():
                fn_name = "screenplay_assets_breakdown"
            fn_args = part.args
        function_call_reply = {
            "functionCall": {
                "name": fn_name,
                "args": fn_args
            }
        }
        return function_call_reply

    def get_reply_text(self, response):
        parts = response.candidates[0].content.parts
        texts = ''
        if parts is None:
            finish_reason = response.candidates[0].finish_reason
            print(f'{finish_reason=}')
            return
        for part in response.candidates[0].content.parts:
            if part.text:
                texts += part.text
            elif part.executable_code:
                continue
        return texts
    
    def get_reply_images(self, context, response):
        parts = response.candidates[0].content.parts
        if parts is None:
            finish_reason = response.candidates[0].finish_reason
            print(f'{finish_reason=}')
            return
        for part in response.candidates[0].content.parts:
            if part.text:
                reply_text = escape(part.text)
                TelegramChannel().send_text(reply_text, context["receiver"])
                logger.info("[TELEGRAMBOT_GEMINI-2.0-FLASH-EXP] sendMsg={}, receiver={}".format(part.text, context["receiver"]))
            elif part.inline_data:
                image = BytesIO(part.inline_data.data)
                logger.info(f"{self.Model_ID} reply={image}")
                image.seek(0)
                TelegramChannel().send_image(image, context["receiver"])
                logger.info("[TELEGRAMBOT_GEMINI-2.0-FLASH-EXP] sendMsg={}, receiver={}".format(image, context["receiver"]))
            elif part.executable_code:
                continue
        return None
    
    def get_search_sources(self, response):
        """
        Get search sources from the response Grounded with Google Search
        """
        sources = []
        ground_chunks = response.candidates[0].grounding_metadata.grounding_chunks
        for i, ground_chunk in enumerate(ground_chunks):
            title = ground_chunk.web.title
            title = escape(title)
            uri = ground_chunk.web.uri
            source = f'{i+1}\.[{title}]({uri})'
            sources.append(source)
        inline_url = '\n'.join(sources)
        return inline_url

    def _convert_to_gemini_1_messages(self, messages: list):
        res = []
        for msg in messages:
            if msg.get("role") == "user":
                role = "user"
            elif msg.get("role") == "assistant":
                role = "model"
            else:
                continue
            res.append({
                "role": role,
                "parts": [{"text": msg.get("content")}]
            })
        return res

    def _convert_to_gemini_15_messages(self, messages: list):
        res = []
        media_parts = []
        user_parts = []
        assistant_parts = []
        for msg in messages:
            msg_role = msg.get('role')
            msg_content = msg.get('content')
            msg_type = type(msg_content).__name__  #识别消息内容的类型
            if msg.get("role") == "user":
                if msg_type == 'File':
                    media_parts.append(msg_content)
                    continue
                if msg_type in ['str','JpegImageFile','dict']:
                    if media_parts != []:
                        media_parts.append(msg_content)
                        user_parts = media_parts
                        parts = user_parts
                        media_parts = []
                    elif media_parts == []:
                        parts = [msg_content]
                role = "user"

            elif msg_role == "assistant":
                assistant_parts = [msg_content]
                parts = assistant_parts
                role = "model"
            else:
                continue
            res.append({
                "role": role,
                "parts": parts
            })
        return res

    def _filter_gemini_1_messages(self, messages: list):
        res = []
        turn = "user"
        for i in range(len(messages) - 1, -1, -1):
            message = messages[i]
            if message.get("role") != turn:
                continue
            res.insert(0, message)
            if turn == "user":
                turn = "assistant"
            elif turn == "assistant":
                turn = "user"
        return res
    
    def _file_cache(self, context):
        memory.USER_FILE_CACHE[context['session_id']] = {
            "path": context.content,
            "msg": context.get("msg")
        }
        logger.info("[{}] file={} is cached for assistant".format(self.Model_ID, context.content))
        return None
    
    def _file_download(self, file_cache):
        msg = file_cache.get("msg")
        path = file_cache.get("path")
        msg.prepare()
        logger.info("[{}] file={} is downloaded locally".format(self.Model_ID, path))
        return None

    def screenplay_scenes_breakdown(
            self, 
            context, 
            fn_name,
            *, 
            screenplay_title: str = None, 
            total_pages: str = None, 
            total_words: str = None, 
            scenes_list: list = []
        ):
        screenplay_title_no_quotes = screenplay_title.replace('《','').replace('》','')
        # 遍历./tmp目录下的文件,如果文件名中含有screenplay_title,则将其路径赋值给path
        for root, dirs, files in os.walk('./tmp'):
            for file in files:
                if screenplay_title_no_quotes in file and file.endswith(('.docx', '.pdf')):
                    path = os.path.join(root, file)
                    break

        if path[-5:] == '.docx':
            reader = docx.Document(path)
            texts = ''
            line_list = []
            num_id = 1
            # 遍历文档中的段落
            for paragraph in reader.paragraphs:
                # 提取文本内容
                scene_normal = paragraph.text.strip()
                # 检查段落是否有序号
                if paragraph._p.pPr and paragraph._p.pPr.numPr:           
                    scene_normal = f"{num_id}. " + scene_normal
                    num_id += 1
                texts = texts + scene_normal + '\n'
                line_list.append(scene_normal)
            total_pages = len(texts)//500 + 1
        elif path[-4:] == '.pdf':
            reader = PdfReader(path)
            total_pages = len(reader.pages)
            texts = ''.join([page.extract_text() for page in reader.pages])
            # 删除脚标
            footer_pattern = rf"\[{re.escape(screenplay_title_no_quotes)}\]\s*◁[\s\x00-\x1f]*▷\s*\n?"
            texts = re.sub(footer_pattern, "", texts, flags=re.DOTALL)
            line_list = re.split(r'[。！？\n]|\s{2,}', texts)
        total_words = len(texts)
        words_per_page = total_words / total_pages
        # 统计每一场的字数
        paragraph = ""
        sc_count = 0
        counter_dict ={}
        # 定义场次描述规则 
        pattern = r"第.*场|场景.*|\d+\..*"
        for i, v in enumerate(line_list):
            # 只要每行的前3～7个字，符合场次描述规则
            if any(re.match(pattern, v.strip()[:n]) is not None for n in (2, 3, 4, 5, 6, 7, 9, 10)):
                counter_dict[f"scene{sc_count}"] = f'{len(paragraph)}'
                sc_count += 1
                paragraph = ""
            paragraph += v.strip()
        # 循环结束后，捕获最后一场戏的字数
        counter_dict[f"scene{sc_count}"] = f'{len(paragraph)}'
        del counter_dict["scene0"]
        
        for i, v in enumerate(scenes_list):
            scene_id = v['scene_id']
            try:
                words_per_scene = int(counter_dict.get(f"scene{scene_id}"))
            except Exception as e:
                logger.error(f"[TELEGRAMBOT_{self.Model_ID}] scene_id={scene_id} not found in counter_dict, error={e}")
                continue
            pages_per_scene = round(words_per_scene / words_per_page, 2)
            estimated_duration = round(pages_per_scene * 1.2, 2)
            v.update(words = words_per_scene)
            v.update(pages = pages_per_scene)
            v.update(estimated_duration = estimated_duration)
        
        # Save the scenes list to an Excel file
        scenes_list_str = json.dumps(scenes_list, ensure_ascii=False)
        df_scenses_list = pd.read_json(scenes_list_str)
        new_cols = ['scene_id', 'location', 'daynight', 'envirement', 'words', 'pages', 'estimated_duration', 'assets_id']
        df_scenses_list = df_scenses_list.reindex(columns=new_cols)
        file_path = TmpDir().path()+ f"{screenplay_title}_scenes_breakdown.xlsx"
        df_scenses_list.to_excel(
            file_path,
            sheet_name=f'{screenplay_title}_scenes_breakdown', 
            index=False
        )
        logger.info(f"[TELEGRAMBOT_{self.Model_ID}] {file_path} is saved")

        # Send the file to the user
        """file = open(file_path, 'rb')
        TelegramChannel().send_file(file, context["receiver"])
        file.close()
        logger.info("[TELEGRAMBOT_{}] sendMsg={}, receiver={}".format(self.Model_ID, file_path, context["receiver"]))"""

        api_response = {
            'total_pages':total_pages,
            'total_words':total_words,
            'scenes_list':scenes_list
        }
        
        # Create a function response part
        function_response_part = []
        function_response_obj = Part.from_function_response(
            name=fn_name,
            response=api_response,
        )
        function_response_dic = {
            "functionResponse": {
                "name": fn_name,
                "response": {
                    "name": fn_name,
                    "content": api_response
                }
            }
        }
        function_response_part.append(function_response_obj)
        # Create a function response text part
        function_response_comment = "这是函数返回的场景表，总结后简单回复即可。不用包含总字数和详细的场景表内容"
        function_response_text = Part.from_text(
            text=function_response_comment
        )
        function_response_part.append(function_response_text)
        # 将function_response_part转换为字符串
        function_response_str = json.dumps(function_response_dic) + '\n' + function_response_comment 
        return function_response_part, function_response_str
    
    def screenplay_assets_breakdown(self, context, fn_name, *, screenplay_title: str = None, assets_list: list = []):
        for i, v in enumerate(assets_list):
            ref_id = i+1
            v.update(ref_url = f'RefURL_{ref_id}')
        
        assets_list_str = json.dumps(assets_list, ensure_ascii=False)
        df_assets_list = pd.read_json(assets_list_str)
        df_scenes_list = pd.read_excel(f'./tmp/{screenplay_title}_scenes_breakdown.xlsx')
        for i, v in enumerate(df_assets_list['scene_ids']):
            asset_pages = 0
            for n in v:
                # 从scenes_list中取出对应的scene_id所在的行
                scene_row = df_scenes_list[df_scenes_list['scene_id'] == n]
                # 取出该行asset_pages列的值
                asset_page = scene_row['pages'].values[0]
                asset_pages += asset_page
            # 将assets_pages和estimated_asset_duration更新到assets_list中
            df_assets_list.at[i, 'asset_pages'] = asset_pages
            df_assets_list.at[i, 'estimated_asset_duration'] = round(asset_pages * 1.2, 2)

        # Save the assets list to an Excel file
        new_cols = ['asset_type', 'asset_id', 'name', 'scene_ids', 'asset_pages', 'estimated_asset_duration', 'ref_url']
        df_assets_list = df_assets_list.reindex(columns=new_cols)
        assets_breaddown_file_path = TmpDir().path()+ f"{screenplay_title}_assets_breakdown.xlsx"
        df_assets_list.to_excel(
            assets_breaddown_file_path,
            sheet_name=f'{screenplay_title}_assets_breakdown', 
            index=False
        )
        logger.info(f"[TELEGRAMBOT_{self.Model_ID}] {assets_breaddown_file_path} is saved")

        # 统计场景引用的资产,更新scenes_list
        for i, v in enumerate(df_scenes_list['scene_id']):
            assets_id = []
            for asset_id_index, scene_ids in enumerate(df_assets_list['scene_ids']):
                asset_scene_ids = [int(x) for x in scene_ids]
                if v in asset_scene_ids:
                    assets_id.append(df_assets_list['asset_id'][asset_id_index])
            # 将assets_id添加到scenes_list中，转换为字符串形式
            df_scenes_list.at[i, 'assets_id'] = str(assets_id)
        scenes_breaddown_file_path = TmpDir().path()+ f"{screenplay_title}_scenes_breakdown.xlsx"
        df_scenes_list.to_excel(
            scenes_breaddown_file_path,
            sheet_name=f'{screenplay_title}_scenes_breakdown', 
            index=False
        )
        logger.info(f"[TELEGRAMBOT_{self.Model_ID}] {scenes_breaddown_file_path} is saved")

        # Send the scenes_breakdown.xlsx to the user
        with open(scenes_breaddown_file_path, 'rb') as f:
            TelegramChannel().send_file(f, context["receiver"])
        logger.info("[TELEGRAMBOT_{}] sendMsg={}, receiver={}".format(self.Model_ID, scenes_breaddown_file_path, context["receiver"]))

        # Create visualization figure of the assets list
        plt.switch_backend('agg') # Use a non-interactive backend
        plt.rcParams['font.sans-serif'] = ['SimHei'] # Set Chinese font
        plt.rcParams['axes.unicode_minus'] = False # Fix the minus sign display issue
        asset_types = set([asset['asset_type'] for asset in assets_list])
        colormap = {
            'cast': None,
            'extra': 'darkcyan',
            'stunt': 'darkgrey',
            'costume': None,
            'location': 'darkgreen',
            'set': 'darkred',
            'vehicle': 'darkgrey',
            'animal': 'darkgoldenrod',
            'prop': 'darkorange'
        }
        for asset_type in asset_types:
            figsize = (30, 5)
            plt.figure(figsize=figsize)
            subset_assets_list = df_assets_list[df_assets_list['asset_type'] == asset_type]
            ax = subset_assets_list.plot(
                figsize=figsize,
                title=f'{screenplay_title}_Assets_Breakdown-{asset_type.upper()}',  
                kind='bar', 
                x='name', 
                y='estimated_asset_duration', 
                ylabel='资产预计时长(分钟)', 
                color=colormap[asset_type],
                legend=False
            )
            # Add value labels on top of the bars
            for i, v in enumerate(subset_assets_list['estimated_asset_duration']):
                ax.text(i, v, f'{v:.2f}', ha='center', va='bottom', fontsize=10)

            figure_path = TmpDir().path() + f"{screenplay_title}_assets_breakdown-{asset_type.upper()}.png"
            plt.savefig(figure_path, dpi=200, bbox_inches='tight')
            plt.close()
            logger.info(f"[TELEGRAMBOT_{self.Model_ID}] {figure_path} is saved")

            # Send the visualization figure to the user
            with open(figure_path, 'rb') as f:
                f.seek(0)
                TelegramChannel().send_image(f, context["receiver"])
            logger.info("[TELEGRAMBOT_{}] sendMsg={}, receiver={}".format(self.Model_ID, figure_path, context["receiver"]))
        
        api_response = {
            'asset_list':assets_list
        }
 
        # Create a function response part
        function_response_part = []
        function_response_obj = Part.from_function_response(
            name=fn_name,
            response=api_response,
        )
        function_response_dic = {
            "functionResponse": {
                "name": fn_name,
                "response": {
                    "name": fn_name,
                    "content": api_response
                }
            }
        }
        function_response_part.append(function_response_obj)
        # Create a function response text part
        function_response_comment = "这是函数返回的资产表，总结后简单回复即可，不用包含详细的资产表内容"
        function_response_text = Part.from_text(
            text=function_response_comment
        )
        function_response_part.append(function_response_text)
        # 将function_response_part转换为字符串
        function_response_str = json.dumps(function_response_dic) + '\n' + function_response_comment 
        return function_response_part, function_response_str
    
    def gemini_15_media(self, query, context, session: BaiduWenxinSession):
        session_id = context.kwargs["session_id"]
        msg = context.kwargs["msg"]
        media_path = context.content
        logger.info(f"[{self.model}] query with media, path={media_path}")

        # Check if the url is a youtube link
        if 'youtube' in media_path:
            media_file = FileData(file_uri=media_path)
        else:
            type_position = media_path.rfind('.') + 1
            mime_type = media_path[type_position:]
            if mime_type in const.IMAGE:
                type_id = 'image'
                # Image URL request
                if query[:8] == 'https://':
                    media_file = media_path

                # Image base64 encoded request
                else:
                    msg.prepare()
                    img = Image.open(media_path)
                    media_file = img
                    # check if the image has an alpha channel
                    if img.mode in ('RGBA','LA') or (img.mode == 'P' and 'transparency' in img.info):
                        # Convert the image to RGB mode,whick removes the alpha channel
                        img = img.convert('RGB')
                        # Save the converted image
                        img_path_no_alpha = media_path[:len(media_path)-3] + 'jpg'
                        img.save(img_path_no_alpha)
                        # Update img_path with the path to the converted image
                        media_file = img_path_no_alpha
            elif mime_type in const.AUDIO:
                msg.prepare()
                type_id = 'audio'
            elif mime_type in const.VIDEO:
                msg.prepare()
                type_id = 'video'
            elif mime_type in const.DOCUMENT:
                msg.prepare()
                # Read and b64encode the PDF file smaller than 20MB
                if mime_type == 'pdf':
                    with open(media_path, 'rb') as file:
                        pdf_data = file.read()
                        b64 = base64.b64encode(pdf_data).decode('utf-8')
                elif mime_type == 'docx':
                    doc = docx.Document(media_path)
                    full_text = []
                    for paragraph in doc.paragraphs:
                        full_text.append(paragraph.text)
                    docx_text = '\n'.join(full_text)
                    b64 = docx_text
                type_id = 'application'
                media_file = {
                    'mime_type': f'{type_id}/{mime_type}',
                    'data': b64
                }
            elif mime_type in const.SPREADSHEET or mime_type in const.TXT:
                msg.prepare()
                type_id = 'text'
            elif mime_type in const.PRESENTATION:
                msg.prepare()
                type_id = 'application'
                mime_type = 'vnd.google-apps.presentation'
            elif mime_type in const.APPLICATION:
                msg.prepare()
                type_id = 'application'
            # Clear original media file in user content avoiding duplicated commitment
            session.messages.pop()
            if (mime_type not in const.IMAGE) and (mime_type not in const.DOCUMENT):
                media_file = self.upload_to_gemini(media_path, mime_type=f'{type_id}/{mime_type}')
        self.sessions.session_query(media_file, session_id)
        self.cache_media(media_path, media_file, context)
    
    def cache_media(self, media_path, media_file, context):
        session_id = context["session_id"]
        if session_id not in memory.USER_IMAGE_CACHE:
            memory.USER_IMAGE_CACHE[session_id] = {
                "path": [media_path],
                "files": [media_file]
            }
        else:
            memory.USER_IMAGE_CACHE[session_id]["path"].append(media_path)
            memory.USER_IMAGE_CACHE[session_id]["files"].append(media_file)
        logger.info(f"[{self.model}] {media_path} cached to memory")
        return None

    def upload_to_gemini(self, path, mime_type=None):
        """Uploads the given file to Gemini.

        https://ai.google.dev/gemini-api/docs/prompting_with_media
        """
        file = self.client.files.upload(
            file=path,
            config=dict(
                display_name=os.path.basename(path),
                mime_type=mime_type
            )
        )
        
        self.wait_for_files_active(file)
        print(f"Uploaded file '{file.display_name}' as: {file.uri}")
        return file
    
    def wait_for_files_active(self, media_file):
        """Waits for the given files to be active.

        Some files uploaded to the Gemini API need to be processed before they can be
        used as prompt inputs. The status can be seen by querying the file's "state"
        field.

        This implementation uses a simple blocking polling loop. Production code
        should probably employ a more sophisticated approach.
        """
        print("Waiting for file processing...")
        file_name =media_file.name
        file = self.client.files.get(name=file_name)
        while file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(10)
            file = self.client.files.get(name=file_name)
        if file.state.name != "ACTIVE":
            raise Exception(f"File {file.name} failed to process")
        print("...all files ready")
        print()
         