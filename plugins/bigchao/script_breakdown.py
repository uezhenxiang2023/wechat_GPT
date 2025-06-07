"""
BigChao API
@author uezhenxiang2023
@Date 2025/05/13
"""

import os
import re
import json
import time

import pandas as pd
import lark_oapi as lark
from pypdf import PdfReader
from google.genai.types import Part
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.oxml import ns, OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from lark_oapi.api.bitable.v1 import *

from config import conf
from common import memory, const, tool_button
from common.log import logger
from common.tmp_dir import TmpDir, create_user_dir

model = conf().get('model')
model_id = const.GEMINI_2_FLASH_IMAGE_GENERATION.upper() if tool_button.imaging is True else model.upper()
app_id = conf().get('feishu_app_id')
app_secret = conf().get('feishu_app_secret')


class AddPageNumber:
    """在文档右上角添加页码"""
    def __init__(self, header):
        self.paragraph = header.paragraphs[0]
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(self, *, font_name='SimHei', font_size=10):
        # 创建页码元素
        run = self.paragraph.add_run()

        # 创建域开始标记
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
        
        # 创建域代码
        instrText = self.create_element('w:instrText')
        instrText.text = "PAGE"
        
        # 创建域结束标记
        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')
        
        # 添加到run中
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # 设置字体格式
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

class CreateBase():
    """创建BreakDown多维表格"""
    def __init__(self, screenplay_title, session_id, scene_records, asset_records):
        self.screenplay_title = screenplay_title
        self.session_id = session_id
        self.scene_records = scene_records
        self.asset_records = asset_records
        # 构建客户端
        self.client = lark.Client.builder() \
        .app_id(app_id) \
        .app_secret(app_secret) \
        .log_level(lark.LogLevel.INFO) \
        .build()
        self.app_token = self.copy_app()
        copy_status = self.copy_status(self.app_token)
        while copy_status != "success":
            time.sleep(10)
            copy_status = self.copy_status(self.app_token)
        self.scene_table_id = self.list_app_table()[0]
        self.asset_table_id = self.list_app_table()[1]

    def copy_app(self):
        """将BD多维表格模版拷贝到tmp目录"""
        # 创建client
        client = self.client

        # 构造请求对象
        request: CopyAppRequest = CopyAppRequest.builder() \
            .app_token("Uen7bebXNaxe1WscVAec3O3Zn5f") \
            .request_body(CopyAppRequestBody.builder()
                .name(f"{self.screenplay_title}_BD_database_{self.session_id}")
                .folder_token("GO4Kf8RgSlXZzgdhtg8cJw1Tn9d")
                .without_content(False)
                .time_zone("Asia/Shanghai")
                .build()) \
            .build()

        # 发起请求
        response: CopyAppResponse = client.bitable.v1.app.copy(request)

        # 处理失败返回
        if not response.success():
            logger.error(
                f"client.bitable.v1.app.copy failed, code: {response.code}, msg: {response.msg}, log_id: {response.get_log_id()}, resp: \n{json.dumps(json.loads(response.raw.content), indent=4, ensure_ascii=False)}")
            return
        else:
            # 处理业务结果
            logger.info(lark.JSON.marshal(response.data, indent=4))
            return response.data.app.app_token
    
    def copy_status(self, app_token):
        """获取多维表格的拷贝状态"""
        # 创建client
        client = self.client

        # 构造请求对象
        request: GetAppRequest = GetAppRequest.builder() \
            .app_token(app_token) \
            .build()

        # 发起请求
        response: GetAppResponse = client.bitable.v1.app.get(request)

        # 处理失败返回
        if not response.success():
            logger.error(
                f"client.bitable.v1.app.get failed, code: {response.code}, msg: {response.msg}, log_id: {response.get_log_id()}, resp: \n{json.dumps(json.loads(response.raw.content), indent=4, ensure_ascii=False)}")
            return

        else:
            # 处理业务结果
            logger.info(lark.JSON.marshal(response.data, indent=4))
            return response.msg
    
    def list_app_table(self):
        """列出多维表格中的数据表"""
    # 创建client
        client = self.client

        # 构造请求对象
        request: ListAppTableRequest = ListAppTableRequest.builder() \
            .app_token(self.app_token) \
            .page_size(20) \
            .build()

        # 发起请求
        response: ListAppTableResponse = client.bitable.v1.app_table.list(request)

        # 处理失败返回
        if not response.success():
            logger.error(
                f"client.bitable.v1.app_table.list failed, code: {response.code}, msg: {response.msg}, log_id: {response.get_log_id()}, resp: \n{json.dumps(json.loads(response.raw.content), indent=4, ensure_ascii=False)}")
            return
        else:
            # 处理业务结果
            items = response.data.items
            # 提取资产表与场景表的table_id
            for item in items:
                if item.name[-7:] == "assects":
                    assects_bd_table = item.table_id
                elif item.name[-6:] == 'scenes':
                    scenes_bd_table = item.table_id
                    logger.info(lark.JSON.marshal(response.data, indent=4))
            return scenes_bd_table, assects_bd_table
        
    def add_scene_table_record(self):
        """向多维表格的场景数据表中批量添加记录"""
        # 创建client
        client = self.client 

        # 构造请求对象
        request: BatchCreateAppTableRecordRequest = BatchCreateAppTableRecordRequest.builder() \
            .app_token(self.app_token) \
            .table_id(self.scene_table_id) \
            .request_body(BatchCreateAppTableRecordRequestBody.builder()
                .records(self.scene_records)
                .build()) \
            .build()

        # 发起请求
        response: BatchCreateAppTableRecordResponse = client.bitable.v1.app_table_record.batch_create(request)

        # 处理失败返回
        if not response.success():
            logger.error(
                f"client.bitable.v1.app_table_record.batch_create failed, code: {response.code}, msg: {response.msg}, log_id: {response.get_log_id()}, resp: \n{json.dumps(json.loads(response.raw.content), indent=4, ensure_ascii=False)}")
            return
        else:
            # 处理业务结果
            logger.info(lark.JSON.marshal(response.data, indent=4))

    def add_asset_table_record(self):
        """向多维表格的资产数据表中批量添加记录"""
        # 创建client
        client = self.client

        # 构造请求对象
        request: BatchCreateAppTableRecordRequest = BatchCreateAppTableRecordRequest.builder() \
            .app_token(self.app_token) \
            .table_id(self.asset_table_id) \
            .request_body(BatchCreateAppTableRecordRequestBody.builder()
                .records(self.asset_records)
                .build()) \
            .build()

        # 发起请求
        response: BatchCreateAppTableRecordResponse = client.bitable.v1.app_table_record.batch_create(request)

        # 处理失败返回
        if not response.success():
            logger.error(
                f"client.bitable.v1.app_table_record.batch_create failed, code: {response.code}, msg: {response.msg}, log_id: {response.get_log_id()}, resp: \n{json.dumps(json.loads(response.raw.content), indent=4, ensure_ascii=False)}")
            return
        else:
            # 处理业务结果
            logger.info(lark.JSON.marshal(response.data, indent=4))

def scene_records(bd_file_path:str):
    """构建批量场景记录消息体"""
    df_scenes_list = pd.read_excel(bd_file_path)
    # 将scene_id字段转换为字符串类型
    df_scenes_list['scene_id'] = df_scenes_list['scene_id'].astype(str)
    scenes_list = df_scenes_list.to_dict('records')
    for item in scenes_list:
        if 'assets_id' in item:
            # Remove quotes and spaces from the string, then split by comma
            assets_id_str = item['assets_id'].replace("'", "").replace("[", "").replace("]", "").strip()
            # Convert to list, excluding empty strings
            item['assets_id'] = [x.strip() for x in assets_id_str.split(',') if x.strip()]
    scene_records = [AppTableRecord.builder().fields(v).build() for v in scenes_list]
    return scene_records

def asset_records(bd_file_path:str):
    """构建批量资产记录消息体"""
    df_assets_list = pd.read_excel(bd_file_path)
    assets_list = df_assets_list.to_dict('records')
    for item in assets_list:
        if 'visual_effects_type' in item and pd.isna(item['visual_effects_type']):
            item['visual_effects_type'] = ""
        if 'visual_effects_description' in item and pd.isna(item['visual_effects_description']):
            item['visual_effects_description'] = ""
        if 'scene_ids' in item:
            # Remove quotes and spaces from the string, then split by comma
            scene_ids_str = str(item['scene_ids']).replace("'", "").replace("[", "").replace("]", "").strip()
            # Convert to list, excluding empty strings
            item['scene_ids'] = [x.strip() for x in scene_ids_str.split(',') if x.strip()]
    asset_records = [AppTableRecord.builder().fields(v).build() for v in assets_list]
    return asset_records

def screenplay_formatter(
        session_id,
        fn_name: str = None,
        *,
        screenplay_title: str = None,
        screenwriter: str = None,
        paragraphs_metadata: list = []
    ):
    # 创建新文文档
    doc = Document()
    # 页面设置为A4尺寸
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # 设置页边距
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    # 添加页码
    page_number = AddPageNumber(section.header)
    paragraph = page_number.add_page_number(font_size=10)
    # 添加封面
    add_cover(doc, screenplay_title=screenplay_title, screenwriter=screenwriter)
    for paragraph_metadata in paragraphs_metadata:
        category = paragraph_metadata['category']
        if category == 'scene_setting':
            add_scene_setting(doc, paragraph_metadata)
        elif category == 'action':
            add_action(doc, paragraph_metadata)
        elif category == 'character':
            add_character(doc, paragraph_metadata)
        elif category == 'dialogue':
            add_dialogue(doc, paragraph_metadata)
    user_dir = TmpDir().path() + str(session_id) + '/response/'
    user_dir_exists = os.path.exists(user_dir)
    if not user_dir_exists:
        create_user_dir(user_dir)
    docx_file_path = user_dir + f"{screenplay_title}.docx"
    pdf_file_path = user_dir + f"{screenplay_title}.pdf"
    # 将文档存储到本地目录
    doc.save(docx_file_path)
    logger.info(f"[TELEGRAMBOT_{model_id}] {docx_file_path} is saved")
    convert(docx_file_path, pdf_file_path)
    logger.info(f"[TELEGRAMBOT_{model_id}] {pdf_file_path} is saved")
    api_response = {
        'file_pathes': [docx_file_path, pdf_file_path]
    }
    # Create a function response part
    function_response_part = []
    function_response_obj = Part.from_function_response(
        name=fn_name,
        response=api_response,
    )
    function_response_part.append(function_response_obj)
    # Create a function response text part
    function_response_comment = "这是按照好莱坞工业标准，排版后的剧本，两种格式，pdf格式方便阅读，docx方便修改"
    function_response_text = Part.from_text(
        text=function_response_comment
    )
    function_response_part.append(function_response_text)
    return function_response_part

def add_cover(
        doc,
        *, 
        screenplay_title:str = None, 
        screenwriter:str = None
    ):
        """添加封面"""
        # 插入空行达到第8行的效果
        for _ in range(7):
            doc.add_paragraph()
        
        # 插入剧本标题段落
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
        title_paragraph.paragraph_format.line_spacing = 1.5
        title_paragraph.paragraph_format.space_after = Pt(0)
        
        # 添加剧本标题
        title_run = title_paragraph.add_run(screenplay_title)
        title_font = title_run.font
        title_font.name = 'SimHei'  # 黑体
        title_font.size = Pt(22)    # 22号字
        title_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        title_font.bold = True
        
        # 添加编剧信息
        if screenwriter:   
            writer_paragraph = doc.add_paragraph()
            writer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            writer_run = writer_paragraph.add_run(f"编剧：{screenwriter}")
            writer_font = writer_run.font
            writer_font.name = 'SimHei'
            writer_font.size = Pt(12)
            writer_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        
        # 添加分节符，开始新的章节
        doc.add_section()
        
def add_scene_setting(doc, content):
        scene_id = content['scene_heading']['scene_id']
        env = content['scene_heading']['enviroment']
        location = content['scene_heading']['location']
        daynight = content['scene_heading']['daynight']
        content = str(scene_id) + ' ' + env + ' ' + location + ' - ' + daynight
        """添加场景段落"""
        # 插入新段落
        paragraph = doc.add_paragraph()
        # 设置段落行间距
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_after = Pt(0)
        # 设置段落左右缩进
        paragraph.paragraph_format.left_indent = Mm(0)
        paragraph.paragraph_format.right_indent = Mm(0)
        # 设置字体
        run = paragraph.add_run(text=content)
        font = run.font
        font.name = 'SimHei'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')  # 东亚文字字体
        font.size = Pt(12)
        font.bold = True
        # 动作段落结尾插入新的空白段落
        paragraph = doc.add_paragraph()
        # 设置段落行间距
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_after = Pt(0)
        # 设置段落左右缩进
        paragraph.paragraph_format.left_indent = Mm(0)
        paragraph.paragraph_format.right_indent = Mm(0)

def add_action(doc, content):
    """添加动作段落"""
    # 插入新段落
    paragraph = doc.add_paragraph()
    # 设置段落行间距
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)
    # 设置段落左右缩进
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.right_indent = Mm(0)
    # 设置字体
    content = convert_punctuation(content['content'])
    run = paragraph.add_run(text=content)
    font = run.font
    font.name = 'SimHei'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')  # 东亚文字字体
    font.size = Pt(12)
    # 动作段落结尾插入新的空白段落
    paragraph = doc.add_paragraph()
    # 设置段落行间距
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)
    # 设置段落左右缩进
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.right_indent = Mm(0)

def add_character(doc, content):
    """添加角色段落"""
    # 插入新段落
    paragraph = doc.add_paragraph()
    # 设置段落行间距
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)
    # 设置段落左右缩进
    paragraph.paragraph_format.left_indent = Mm(56)
    paragraph.paragraph_format.right_indent = Mm(0)
    # 检查角色名是否以冒号结尾，如果不是则添加
    character_text = content['content']
    # 替换特殊标记 (不区分大小写)
    character_text = re.sub(r'\bos\b', '(O.S.)', character_text, flags=re.IGNORECASE)
    character_text = re.sub(r'\bvo\b', '(V.O.)', character_text, flags=re.IGNORECASE)
    if not character_text.endswith((':', '：')):
        character_text += '：'  # 使用中文冒号  
    # 设置字体
    run = paragraph.add_run(text=character_text)
    font = run.font
    font.name = 'SimHei'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')  # 东亚文字字体
    font.size = Pt(12)

def add_dialogue(doc, content):
    """添加台词段落"""
    # 插入新段落
    paragraph = doc.add_paragraph()
    # 设置段落行间距
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)
    # 设置段落左右缩进
    paragraph.paragraph_format.left_indent = Mm(30)
    paragraph.paragraph_format.right_indent = Mm(30)
    # 设置字体
    dialogue_content = convert_punctuation(content['content'])
    run = paragraph.add_run(text=dialogue_content)
    font = run.font
    font.name = 'SimHei'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')  # 东亚文字字体
    font.size = Pt(12)
    # 台词段落结尾插入新的空白段落
    paragraph = doc.add_paragraph()
    # 设置空白段落行间距
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)
    # 设置空白段落左右缩进
    paragraph.paragraph_format.left_indent = Mm(30)
    paragraph.paragraph_format.right_indent = Mm(30)

def screenplay_scenes_breakdown(
        session_id,
        fn_name,
        *,
        screenplay_title: str = None,
        total_pages: str = None,
        total_words: str = None,
        scenes_list: list = []
    ):
    """scenes breakdown for screenscript"""
    screenplay_title_no_quotes = screenplay_title.replace('《', '').replace('》', '')
    # 遍历./tmp目录下的文件,如果文件名中含有screenplay_title,则将其路径赋值给path
    for root, dirs, files in os.walk(f'./tmp/{session_id}'):
        for file in files:
            if screenplay_title_no_quotes in file and file.endswith(('.docx', '.pdf')):
                path = os.path.join(root, file)
                break
    if path.endswith('.pdf'):
        total_words, words_per_page, counter_dict = sc_counter(path)
    """elif path.endswith('.docx'):
        reader = docx.Document(path)
        texts = ''
        line_list = []
        num_id = 1
        # 遍历文档中的段落
        for paragraph in reader.paragraphs:
            # 提取文本内容
            scene_normal = paragraph.text.strip()
            # 检查段落是否有序号
            if paragraph._p.pPr and paragraph._p.pPr.numPr:
                scene_normal = f"{num_id}. " + scene_normal
                num_id += 1
            texts = texts + scene_normal + '\n'
            line_list.append(scene_normal)
        total_pages = len(texts)//500 + 1"""
    for i, v in enumerate(scenes_list):
        scene_id = v['scene_id']
        try:
            words_per_scene = int(counter_dict.get(f"scene{scene_id}"))
        except Exception as e:
            logger.error(f"[TELEGRAMBOT_{model_id}] scene_id={scene_id} not found in counter_dict, error={e}")
            continue
        pages_per_scene = round(words_per_scene / words_per_page, 2)
        estimated_duration = round(pages_per_scene * 1.2, 2)
        v.update(words = words_per_scene)
        v.update(pages = pages_per_scene)
        v.update(estimated_duration = estimated_duration)

    # Save the scenes list to an Excel file
    scenes_list_str = json.dumps(scenes_list, ensure_ascii=False)
    df_scenses_list = pd.read_json(scenes_list_str)
    new_cols = ['scene_id', 'location', 'daynight', 'envirement', 'words', 'pages', 'estimated_duration', 'assets_id']
    df_scenses_list = df_scenses_list.reindex(columns=new_cols)
    user_dir = TmpDir().path() + str(session_id) + '/response/'
    user_dir_exists = os.path.exists(user_dir)
    if not user_dir_exists:
        create_user_dir(user_dir)
    file_path = user_dir + f"{screenplay_title}_scenes_breakdown.xlsx"
    df_scenses_list.to_excel(
        file_path,
        sheet_name=f'{screenplay_title}_scenes_breakdown',
        index=False
    )
    logger.info(f"[TELEGRAMBOT_{model_id}] {file_path} is saved")
    api_response = {
        'total_pages': total_pages,
        'total_words': total_words,
        'scenes_list': scenes_list
    }
    # Create a function response part
    function_response_part = []
    function_response_obj = Part.from_function_response(
        name=fn_name,
        response=api_response,
    )
    function_response_part.append(function_response_obj)
    # Create a function response text part
    function_response_comment = "这是函数返回的原始场景表，主要用来统计页数和时长，不用针对该表内容进行回复。"
    function_response_text = Part.from_text(
        text=function_response_comment
    )
    function_response_part.append(function_response_text)
    return function_response_part

def screenplay_assets_breakdown(
        session_id,
        fn_name,
        *,
        screenplay_title: str = None,
        assets_list: list = []
    ):
    """assets breaddown for screenscript"""
    for i, v in enumerate(assets_list):
        ref_id = i+1
        v.update(ref_url = f'RefURL_{ref_id}')

    assets_list_str = json.dumps(assets_list, ensure_ascii=False)
    df_assets_list = pd.read_json(assets_list_str)
    user_dir = TmpDir().path() + str(session_id) + '/response/'
    df_scenes_list = pd.read_excel(user_dir + f'{screenplay_title}_scenes_breakdown.xlsx')
    for i, v in enumerate(df_assets_list['scene_ids']):
        asset_pages = 0
        for n in v:
            # 从scenes_list中取出对应的scene_id所在的行
            scene_row = df_scenes_list[df_scenes_list['scene_id'] == n]
            # 取出该行asset_pages列的值
            asset_page = scene_row['pages'].values[0]
            asset_pages += asset_page
        # 将assets_pages和estimated_asset_duration更新到assets_list中
        df_assets_list.at[i, 'asset_pages'] = asset_pages
        df_assets_list.at[i, 'estimated_asset_duration'] = round(asset_pages * 1.2, 2)

    # Save the assets list to an Excel file
    new_cols = ['asset_type', 'asset_id', 'name', "visual_effects_type", 'visual_effects_description', 'scene_ids', 'asset_pages', 'estimated_asset_duration', 'ref_url']
    df_assets_list = df_assets_list.reindex(columns=new_cols)
    assets_list = df_assets_list.to_dict('records')
    assets_breakdown_file_path = user_dir + f"{screenplay_title}_assets_breakdown.xlsx"
    df_assets_list.to_excel(
        assets_breakdown_file_path,
        sheet_name=f'{screenplay_title}_assets_breakdown',
        index=False
    )
    logger.info(f"[TELEGRAMBOT_{model_id}] {assets_breakdown_file_path} is saved")
    # 为资产构建飞书多维表格的批量记录
    asset_records_list = asset_records(assets_breakdown_file_path)

    # 统计场景引用的资产,更新scenes_lists
    for i, v in enumerate(df_scenes_list['scene_id']):
        assets_id = []
        for asset_id_index, scene_ids in enumerate(df_assets_list['scene_ids']):
            asset_scene_ids = [int(x) for x in scene_ids]
            if v in asset_scene_ids:
                assets_id.append(df_assets_list['asset_id'][asset_id_index])
        # 将assets_id添加到scenes_list中，转换为字符串形式
        df_scenes_list.at[i, 'assets_id'] = str(assets_id)
    scenes_list = df_scenes_list.to_dict('records')
    scenes_breakdown_file_path = user_dir + f"{screenplay_title}_scenes_breakdown.xlsx"
    df_scenes_list.to_excel(
        scenes_breakdown_file_path,
        sheet_name=f'{screenplay_title}_scenes_breakdown',
        index=False
    )
    logger.info(f"[TELEGRAMBOT_{model_id}] {scenes_breakdown_file_path} is updated")
    # 为场景构建飞书多维表格的批量记录
    scene_records_list = scene_records(scenes_breakdown_file_path)
    # 创建多维表格
    base = CreateBase(screenplay_title, session_id, scene_records_list, asset_records_list)
    base.add_scene_table_record()
    base.add_asset_table_record()
    logger.info(f"[TELEGRAMBOT_{model_id}][Lark_Base]{screenplay_title} is created")


    """# Send the scenes_breakdown.xlsx to the user
    with open(scenes_breakdown_file_path, 'rb') as f:
        TelegramChannel().send_file(f, context["receiver"])
    logger.info("[TELEGRAMBOT_{}] sendMsg={}, receiver={}".format(self.Model_ID, scenes_breakdown_file_path, context["receiver"]))"""

    """# Create visualization figure of the assets list
    plt.switch_backend('agg') # Use a non-interactive backend
    plt.rcParams['font.sans-serif'] = ['SimHei'] # Set Chinese font
    plt.rcParams['axes.unicode_minus'] = False # Fix the minus sign display issue
    asset_types = set([asset['asset_type'] for asset in assets_list])
    colormap = {
        'cast': None,
        'extra_silent': 'darkcyan',
        'extra_atmosphere': 'darkblue',
        'stunts': 'darkgrey',
        'costume': None,
        'makeup': None,
        'location': 'darkgreen',
        'set_dressing': 'darkred',
        'vehicle': 'darkgrey',
        'animal': 'darkgoldenrod',
        'livestock': None,
        'prop': 'darkorange',
        'special_effects': None
    }
    for asset_type in asset_types:
        figsize = (30, 5)
        plt.figure(figsize=figsize)
        subset_assets_list = df_assets_list[df_assets_list['asset_type'] == asset_type]
        ax = subset_assets_list.plot(
            figsize=figsize,
            title=f'{screenplay_title}_Assets_Breakdown-{asset_type.upper()}',  
            kind='bar', 
            x='name', 
            y='estimated_asset_duration', 
            ylabel='资产预计时长(分钟)', 
            color=colormap[asset_type],
            legend=False
        )
        # Add value labels on top of the bars
        for i, v in enumerate(subset_assets_list['estimated_asset_duration']):
            ax.text(i, v, f'{v:.2f}', ha='center', va='bottom', fontsize=10)

        figure_path = TmpDir().path() + f"{screenplay_title}_assets_breakdown-{asset_type.upper()}.png"
        plt.savefig(figure_path, dpi=200, bbox_inches='tight')
        plt.close()
        logger.info(f"[TELEGRAMBOT_{self.Model_ID}] {figure_path} is saved")

        # Send the visualization figure to the user
        with open(figure_path, 'rb') as f:
            f.seek(0)
            TelegramChannel().send_image(f, context["receiver"])
        logger.info("[TELEGRAMBOT_{}] sendMsg={}, receiver={}".format(self.Model_ID, figure_path, context["receiver"]))"""

    api_response = {
        'scenes_list': json.dumps(scenes_list),
        'assets_list': json.dumps(assets_list),
        'file_pathes': [scenes_breakdown_file_path, assets_breakdown_file_path]
    }
    # Create a function response part
    function_response_part = []
    function_response_obj = Part.from_function_response(
        name=fn_name,
        response=api_response,
    )
    function_response_part.append(function_response_obj)
    # Create a function response text part
    function_response_comment = "这是函数返回的最终场景表和资产表，回复时不用包含表中详细内容，总结后简单回复即可。"
    function_response_text = Part.from_text(
        text=function_response_comment
    )
    function_response_part.append(function_response_text)
    return function_response_part

def cache_media(media_path, media_file, context):
        session_id = context["session_id"]
        if session_id not in memory.USER_IMAGE_CACHE:
            memory.USER_IMAGE_CACHE[session_id] = {
                "path": [media_path],
                "files": [media_file]
            }
        else:
            memory.USER_IMAGE_CACHE[session_id]["path"].append(media_path)
            memory.USER_IMAGE_CACHE[session_id]["files"].append(media_file)
        logger.info(f"[{model_id}] {media_path} cached to memory")
        return None

def convert_punctuation(text: str) -> str:
    """将英文标点转换为中文标点"""
    punctuation_map = {
        ',': '，',
    }
    
    for en_punct, cn_punct in punctuation_map.items():
        text = text.replace(en_punct, cn_punct)
    return text

def sc_counter(pdf_path:str) -> dict:
    """精算每一场的字数.

    args:
        pdf_path: 剧本的本地存储路径，如'./tmp/《炒房客》_v2.pdf'
    returns:
        total_words： 总字数\n
        words_per_page: 页均字数\n
        counter_dict: 场号及对应场次内容的字数.
    """
    reader = PdfReader(pdf_path)
    total_pages = len(reader.pages)
    texts = ''.join([page.extract_text() for page in reader.pages])
    # 1. 清理纯数字页码（位于每行开头或结尾的数字）
    texts = re.sub(r'\s*\d+\s*$', '', texts, flags=re.MULTILINE)
    # 2. 删除换行符
    texts = texts.replace('\t', ' ').strip()
    # 3. 删除脚标
    footer_pattern = rf"\[{re.escape('炒房客')}\]\s*◁[\s\x00-\x1f]*▷\s*\n?"
    texts = re.sub(footer_pattern, "", texts, flags=re.DOTALL)
    total_words = len(texts)
    words_per_page = total_words / total_pages
    # 用句号、感叹号、问号和换行符拆分texts内容
    lines_list = re.split(r'[。！？\n]|\s{2,}', texts)
    # 统计每一场的字数
    paragraph = ""
    sc_heading = ""
    sc_id = 0
    counter_dict ={}
    # 定义场次描述规则
    pattern = r"第.*场|场景.*|\d+\..*|\d+.*"
    for i, v in enumerate(lines_list):
        # 只要每行的前3～7个字，符合场次描述规则
        if any(re.match(pattern, v.strip()[:n]) is not None for n in (2, 3, 4, 5, 6, 7, 9, 10)):
            counter_dict[f"scene{sc_id}"] = f'{len(paragraph)-len(sc_heading)}'
            sc_id += 1
            sc_heading = v
            paragraph = ""
        paragraph += v.strip()
    # 循环结束后，捕获最后一场戏的字数
    counter_dict[f"scene{sc_id}"] = f'{len(paragraph)}'
    del counter_dict["scene0"]
    return total_words, words_per_page, counter_dict